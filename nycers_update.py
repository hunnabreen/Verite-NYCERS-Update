#import docx
from docx import *
from docx.enum.text import WD_COLOR_INDEX
#import os for splitting file types
import os


def main():
    print("Before starting, ensure that all documents that need updating are in the same folder as this program")
    #init original document
    file = input("Enter file here: ")
    document = Document(file)
    """
    open(document)
    line = document.readlines
    print(line)
        
    #add every paragraph to a map, 
    #key = paragraph number 
    #value = every run in paragraph as nested key, style for run as nested value
    #maybe have a nested map? Map<para, map<runs,style>>
   
    para = {} # paragraph key
    def runLoader(document):
        for p in document.paragraphs:
            #add paragraph no. to key
            vals = {} #nested key
            para[p] = vals
            for r in p.runs:
                #add runs & styles to vals mapper
                vals[r] = r.style
                if next(r) == None:
                    break
        return

    
        parses the document to update highlights
        red => purple
        green => blue
        purple => delete
        blue => unhighlight
    """
    #edit so it can access key-value pairs in para map rather than just the doc
    def textUpdate(document):
        doc = document
        for chars in doc.paragraphs:
            for run in chars.runs:
                highlight = run.font.highlight_color
                if run.font.highlight_color == WD_COLOR_INDEX.RED:
                    run.font.highlight_color = WD_COLOR_INDEX.PINK
                    
                elif run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    
                elif run.font.highlight_color == WD_COLOR_INDEX.PINK:
                    run.clear() #delete
                    
                elif run.font.highlight_color == WD_COLOR_INDEX.TURQUOISE:
                    run.font.highlight_color = WD_COLOR_INDEX.WHITE #unhighlight
                    
                else:
                    break

    #updates the bullet points in the word docx
    def bulletUpdate(document):
        doc = document
        for para in doc.paragraphs:
            sty = para.style
            if sty == 'List Bullet':
                if sty.highlight_color == WD_COLOR_INDEX.RED:
                    sty.highlight_color =  WD_COLOR_INDEX.PINK
                else:
                    break
            else:
                break
      
    
    #assemble the updated docx back together
    def assemble(para):
        #take in para and assemble it back together on new document
        return

    #driver
    def driver(document):
        #runLoader(document)
        textUpdate(document)
        bulletUpdate(document)
        #assemble(document)
        #remove .docx file extension and save as UPDATED.docx
        #file = os.path.splitext(file)[0] + '-UPDATE.docx'
        document.save(file + '-UPDATE.docx')
        print("document updated, look at the file directory for updated file")
    
    driver(document)    

main()
